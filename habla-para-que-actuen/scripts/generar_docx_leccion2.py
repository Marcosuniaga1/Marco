#!/usr/bin/env python3
"""Genera el .docx de la Lección 2 — El bloqueo tiene nombre"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

GUION = """En la lección anterior grabaste un audio de sesenta segundos. Identificaste dónde dudaste, dónde aceleraste, dónde no sonaste como querías. Y escribiste una frase: "Mi mayor problema al comunicar es..." con tus propias palabras.

Hoy vamos a trabajar con esa frase. Porque lo que escribiste ahí no es casual. Tiene una mecánica detrás. Y si no la entiendes, cualquier técnica que yo te dé va a funcionar un par de días y después tu cerebro va a volver a hacer lo de siempre.

Te cuento algo que me pasó a mí. Yo llevaba años en televisión cuando me tocó hacer mi primera presentación como formador, ya no como actor. Era un grupo de quince empresarios en una sala en Caracas. Yo sabía el contenido perfectamente, lo había preparado durante semanas. Y tres segundos antes de empezar, mi cabeza me dijo: "Estos tipos manejan empresas. ¿Qué les vas a enseñar tú, que vienes de hacer telenovelas?"

Tres segundos. Eso fue lo que tardó mi cerebro en desmontar semanas de preparación. Empecé a hablar y lo que salió fue una versión diluida de lo que había preparado. Hablé más rápido, me salté ejemplos, y terminé veinte minutos antes de lo previsto. Después, en el carro, me pregunté qué había pasado. Yo sabía el material. Tenía experiencia hablando frente a cámaras, frente a equipos de producción, frente a directores exigentes. Pero esa situación específica activó algo diferente.

Lo que me pasó tiene una explicación concreta. En los segundos antes de hablar, tu amígdala, la parte del cerebro que evalúa amenazas, escanea la situación y decide si estás en peligro. Y tu amígdala no distingue entre un león y una cámara. Para ella, una amenaza social activa la misma señal que una amenaza física.

Cuando eso pasa, tu cuerpo responde. El cortisol sube. Es la hormona del estrés. En dosis pequeñas te mantiene alerta, en dosis altas te paraliza. Cuando sientes que se te seca la boca antes de hablar, eso es cortisol. Cuando la idea que tenías clara de repente se vuelve confusa, eso es cortisol limitando tu acceso a la parte del cerebro donde organizas ideas y armas frases coherentes.

Tu respiración se acorta. Tu cuerpo se prepara para correr o pelear, y la consecuencia directa es que tu voz pierde potencia. Te sale débil, entrecortada, o demasiado rápida porque estás metiendo muchas palabras en poca exhalación.

Y tu postura se cierra. Los hombros hacia adelante, los brazos pegados al cuerpo, la mirada baja. Tu cuerpo dice quiero ocupar menos espacio, que no me vean. Exactamente lo contrario de lo que necesitas comunicar.

Todo eso pasa en tres segundos. Y la mayoría de la gente solo sabe que le dan nervios.

Mira, nervios no es un diagnóstico. Es como ir al médico y decirle me duele. ¿Dónde te duele? ¿Desde cuándo? ¿Qué lo activa? Sin esas respuestas no se puede hacer nada. Lo que necesitas saber es qué tipo de situación específica activa tu respuesta, qué pensamiento concreto aparece en tu cabeza, y qué reacción física se dispara primero. Cuando tienes esas tres respuestas, tienes lo que en programación neurolingüística se llama un disparador. Y un disparador sí se puede trabajar.

En mi experiencia hay cuatro disparadores que se repiten constantemente. Voy a describirlos y quiero que identifiques cuál es el tuyo. Puedes tener más de uno, pero siempre hay uno que manda.

El primero es el juicio. Suena así en tu cabeza: van a pensar que no sé de lo que hablo. Algunos le llaman síndrome del impostor, pero a mí me interesa la mecánica. ¿Por qué crees que te van a juzgar? Normalmente porque te estás comparando con alguien que percibes como superior. Un competidor, un referente, alguien con más seguidores. Y esa comparación la haces justo antes de hablar. No tres días antes. Tres segundos antes. Eso fue exactamente lo que me pasó en aquella presentación con los empresarios.

El segundo es el perfeccionismo. Si no lo digo perfecto, mejor no lo digo. Común en emprendedores que saben mucho. Quieren decirlo todo, y el miedo a dejar algo fuera o a ser imprecisos los paraliza. Terminan diciendo demasiado, de forma desordenada, y el mensaje se diluye. O directamente no dicen nada porque nada les parece suficientemente bien formulado.

El tercero es la exposición. No me gusta ser el centro de atención. Este no tiene que ver con lo que dices sino con ser visible. La incomodidad de que la gente te mire, de que tu cara quede grabada. Funciona diferente a los otros porque no se activa cuando hablas, se activa antes, cuando piensas en que vas a hablar. Muchas veces el resultado es evitación directa. Simplemente no grabas el video, no subes el contenido, no haces la presentación.

Un emprendedor con el que trabajé hace unos meses tenía exactamente este disparador. Tipo brillante, sabía de su tema más que la mayoría. Pero llevaba ocho meses con un canal de YouTube con cero videos publicados. Los grababa, los revisaba y los borraba. Cuando le pregunté qué pensaba justo antes de borrar, me dijo: "Pienso que alguien que me conoce lo va a ver y va a pensar que me estoy creyendo algo." Eso es exposición pura. El contenido estaba bien. Lo que le bloqueaba era existir en público.

El cuarto es el resultado. ¿Y si hablo y no pasa nada? Este es el más sutil. Haces un video y nadie comenta. Mandas un audio de ventas y no te contestan. Haces una presentación y nadie compra. Tu cerebro empieza a asociar comunicar con esfuerzo sin recompensa. Y para protegerte de esa frustración, te pone trabas antes de que empieces.

Fíjate cómo funciona el patrón completo. Primero viene un detonante: algo te dice que tienes que comunicar. Un pensamiento, una situación, una fecha. Después aparece el pensamiento automático. Tu disparador se activa y genera una frase en tu cabeza que tú no elegiste. Y finalmente la respuesta: el cortisol sube, la respiración se acorta, la postura se cierra, o directamente evitas la situación. Y si no la evitas y hablas, hablas desde ese estado. Tu audiencia lo percibe. No conscientemente. Pero sienten que algo no cuadra entre lo que dices y cómo lo dices. Y la confianza baja.

Ese patrón es automático. No lo puedes detener pensando voy a dejar de ponerme nervioso. Lo que sí puedes hacer es intervenir en un punto específico para interrumpirlo.

La mayoría de los consejos que ves por ahí intentan intervenir en la respuesta. Respira profundo. Piensa en positivo. Haz la pose de poder. Esas cosas ayudan un poco en el momento, pero al día siguiente el patrón sigue ahí intacto.

La intervención real está en el pensamiento automático. En programación neurolingüística se llama reencuadre. No puedes controlar que el pensamiento aparezca, pero puedes cambiar lo que haces con él cuando aparece.

Te lo explico con mi ejemplo. Mi disparador era el juicio. Mi pensamiento automático era ¿qué les vas a enseñar tú, que vienes de telenovelas? El reencuadre que aprendí a hacer no fue decirme sí puedo, soy genial. Eso es afirmación positiva y dura cinco minutos. El reencuadre fue: ese pensamiento apareció, lo reconozco, pero no necesito obedecerlo para empezar a hablar.

Parece simple. Cambia todo. Porque lo que mantiene el patrón vivo no es el pensamiento en sí. Es la creencia de que tiene autoridad. Que si tu cabeza dice vas a hacer el ridículo, eso tiene peso real sobre lo que haces después. La técnica de reencuadre no discute con el pensamiento. Lo observa y decide no actuar según él.

En la lección cinco vamos a integrar esto con un protocolo completo. Pero necesitaba que entendieras la mecánica ahora porque las lecciones tres y cuatro construyen directamente sobre ella.

Ahora viene tu trabajo. Te pido que lo hagas con honestidad.

Haz una lista de las últimas cinco veces que evitaste comunicar algo importante. Que cancelaste un video, que no levantaste la mano en una reunión, que postergaste una llamada de ventas, que empezaste a grabar y borraste. Cinco situaciones concretas.

Para cada una, anota qué pensaste justo antes de decidir no hacerlo. No la razón que te diste después. Lo que realmente pasó por tu cabeza en ese instante.

Busca el patrón. ¿Juicio, perfeccionismo, exposición, resultado? Puede que tengas dos que compiten. Pero uno manda.

Escribe tu disparador principal en una frase: "Me bloqueo cuando..." y complétala. Si te quedó vaga, tipo me bloqueo cuando estoy nervioso, especifica más. "Me bloqueo cuando tengo que hablar frente a personas que considero más exitosas que yo." Eso sí es un disparador.

Guarda esa frase junto a la de la lección uno. Ahora tienes dos piezas de diagnóstico: dónde pierdes impacto y qué lo activa. En la lección tres vamos a pasar de entender a actuar. Vas a aprender herramientas físicas para cambiar cómo te perciben antes de decir una sola palabra. Herramientas que yo usé durante años en televisión y en cine. Que funcionan incluso cuando tienes miedo.

Con eso listo, nos vemos en la lección tres.

Soy Marco Suniaga, y esto es Habla Para Que Actúen."""


def generar_docx(output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    titulo = doc.add_heading("El bloqueo tiene nombre", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Habla Para Que Actúen  |  Lección 2")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    palabras = len(GUION.split())
    duracion = round(palabras / 150)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pairs = [
        ("Duración: ", f"~{duracion} min"),
        ("   Palabras: ", str(palabras)),
        ("   Fecha: ", datetime.now().strftime("%d-%m-%Y")),
        ("   Versión: ", "Humanizada V2"),
    ]
    for label, value in pairs:
        r_label = info.add_run(label)
        r_label.font.size = Pt(10)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(100, 100, 100)
        r_value = info.add_run(value)
        r_value.font.size = Pt(10)
        r_value.font.color.rgb = RGBColor(100, 100, 100)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)

    doc.add_paragraph()

    paragraphs = [p.strip() for p in GUION.split("\n\n") if p.strip()]
    for parrafo in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(40, 40, 40)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep2 = sep2.add_run("─" * 60)
    run_sep2.font.color.rgb = RGBColor(200, 200, 200)
    run_sep2.font.size = Pt(8)
    doc.add_paragraph()

    notas_titulo = doc.add_heading("Notas de grabación", level=2)
    for run in notas_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    notas = [
        "Energía: Intensa pero contenida. Más íntima que la Lección 1.",
        "Ritmo: Pausar después de revelaciones personales (la anécdota de Caracas).",
        "Anécdota personal: Bajar un poco el tono, como si recordara. No dramatizar.",
        "Los 4 disparadores: Cada uno con voz ligeramente diferente al citar el pensamiento interno.",
        "Historia del emprendedor de YouTube: Contar como si fuera reciente, con detalle.",
        "Cierre: Volver a energía alta para el ejercicio y la llamada a acción.",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nota)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    anot_titulo = doc.add_heading("Espacio para anotaciones", level=2)
    for run in anot_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    for _ in range(5):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Validado por TrueGuiones  |  Listo para grabar")
    run_footer.font.size = Pt(9)
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return palabras


if __name__ == "__main__":
    out = "/tmp/claude-0/-home-user-Marco/4317bd10-8d2d-5470-84b8-760b686bf996/scratchpad/Leccion-2-El-bloqueo-tiene-nombre.docx"
    palabras = generar_docx(out)
    print(f"Generado: {out}")
    print(f"Palabras: {palabras}")
