#!/usr/bin/env python3
"""Genera el .docx de la Lección 4 — El orden en que dices las cosas lo cambia todo"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

GUION = """A esta altura ya sabes tres cosas que antes no sabías: dónde pierdes impacto cuando comunicas, por qué tu cerebro te sabotea, y cómo usar tu voz y tu cuerpo para que te perciban diferente. Hoy vamos a trabajar en lo que dices. No en cómo lo dices. En qué dices y en qué orden lo dices.

Porque la presencia sin mensaje es carisma vacío. Y el mensaje sin estructura es información desordenada. Tu audiencia escucha, asiente, y a los cinco minutos no recuerda nada de lo que dijiste.

Lo que te voy a enseñar viene del neuromarketing. Que no es otra cosa que entender cómo tu audiencia procesa la información cuando tiene que tomar una decisión. Y quiero ser directo con algo: esto no es manipulación. Manipulación es hacer que alguien tome una decisión que le perjudica sin que se dé cuenta. Lo que vamos a hacer es lo contrario. Vas a aprender a presentar tu oferta en un orden que el cerebro de tu audiencia pueda procesar correctamente. Porque muchas veces el problema no es tu oferta. Es el orden en que la presentas.

Te lo explico con algo que me pasó cuando empecé a dar formaciones. Yo preparaba mis presentaciones con toda la información ordenada lógicamente. Primero mi historia, después mi método, después los resultados, al final la oferta. Parecía perfecto. Y la gente escuchaba, tomaba notas, me decía que estuvo muy bien. Pero nadie compraba. Después de tres presentaciones así, un colega que venía del mundo de las ventas me dijo algo que me cambió la perspectiva: "Marco, tu presentación es excelente para alguien que ya decidió comprarte. Pero la gente que te está escuchando todavía no decidió nada. Y los estás perdiendo en los primeros treinta segundos."

Tenía razón. Yo empezaba hablando de mí. Y al cerebro de mi audiencia no le importaba quién era yo. Todavía. Le iba a importar después, pero en los primeros segundos estaba procesando una sola pregunta: ¿esto tiene que ver conmigo?

Si la respuesta es sí, prestan atención. Si no pueden determinar la respuesta porque les estás dando información que no contesta esa pregunta, se desconectan. Y la desconexión es rápida. En un video tienes entre cinco y siete segundos. En una presentación tal vez diez. En una llamada de ventas, quince si tienes suerte.

Para entender por qué pasa esto, necesitas saber algo básico sobre cómo funciona el cerebro cuando recibe un mensaje. Lo voy a simplificar mucho porque no necesitas un título en neurociencia para aplicarlo.

Tu audiencia procesa tu mensaje con tres sistemas, en este orden. El primero es el cerebro reptiliano, la parte más antigua. Solo procesa supervivencia. Peligro o seguridad. Es binario, es rápido, y decide si presta atención o no. Cuando alguien dice algo que activa esa parte, algo que suena como un peligro o una oportunidad directa, no puedes ignorarlo aunque quieras.

El segundo es el cerebro emocional, el límbico. Procesa sentimientos, recuerdos, conexiones personales. Cuando sientes que alguien entiende tu situación, eso es tu cerebro límbico activándose. Y sin esa conexión, no hay confianza. Y sin confianza no hay compra. Nunca.

El tercero es el cerebro racional, el neocórtex. Procesa lógica, datos, comparaciones. Y aquí viene lo que lo cambia todo: el cerebro racional no toma decisiones. Justifica decisiones que ya fueron tomadas por los otros dos. Cuando alguien dice "lo pensé bien y decidí que sí", lo que realmente pasó fue que sintió que sí, quiso que sí, y después armó una explicación lógica para justificarlo.

Fíjate lo que significa eso para tu comunicación. Si empiezas con datos y lógica, le estás hablando al tercer cerebro. Al que no decide. Si empiezas activando atención y emoción, capturas los dos primeros y el tercero simplemente confirma.

Por eso te voy a dar un framework de tres bloques. Cada uno apunta a uno de los tres cerebros, en el orden correcto.

El primer bloque es la apertura. Siete segundos. Tu único objetivo es que el cerebro reptiliano diga esto tiene que ver conmigo. Tienes tres herramientas para lograrlo.

La pregunta directa. Una pregunta que describe una situación que tu audiencia vive. No una pregunta retórica vacía. Una que si la leen, no pueden evitar responder mentalmente. Por ejemplo: "¿Cuántas veces esta semana evitaste hacer un video porque no sabías qué decir?" Eso funciona porque es específico, es una acción concreta y apunta a un motivo real.

La afirmación que contradice. Una frase que dice lo contrario de lo que tu audiencia asume. "Tu producto no es el problema. Nunca lo fue." Si tu audiencia lleva meses pensando que necesita mejorar su producto para vender más, esa frase los frena. Y cuando algo contradice lo que crees, tu cerebro necesita resolverlo.

El dato con peso. Un número que cambia la perspectiva. "El noventa y tres por ciento de cómo te perciben no tiene nada que ver con lo que dices." Ese número va contra la intuición y genera curiosidad inmediata.

Lo que no funciona como apertura: "Hola, mi nombre es equis y hoy te voy a hablar de..." Eso es un saludo. No activa nada.

El segundo bloque es el desarrollo. Cuarenta a cincuenta segundos. Aquí le hablas al cerebro emocional. Tiene tres partes.

Primero describes el problema con las palabras de tu audiencia. No "comunicación ineficaz", eso es jerga. Sino "sientes que lo que dices no convence a nadie" o "grabas un video, lo ves y piensas eso no suena como yo quiero sonar". Las palabras exactas las sacas de las conversaciones con tus clientes, de los mensajes que te mandan, de los comentarios en redes.

Después muestras cómo lo resuelves, pero sin dar la solución completa. Aquí presentas tu enfoque, no tu producto. Decir "tengo un curso de cinco lecciones con videos y ejercicios" es describir un producto. Eso le habla al cerebro racional y todavía no es su turno. Decir "hay un sistema de tres capas que cambia tu comunicación de informativa a magnética" es describir un enfoque. Eso le habla al cerebro emocional porque plantea una transformación.

Y finalmente das un resultado concreto. No un testimonio fabricado. Un caso real. "Un emprendedor que aplicó esto pasó de evitar las cámaras a grabar tres videos por semana y duplicar sus llamadas entrantes en un mes." Eso tiene un antes, un después y un número. El cerebro no necesita que le digas "funciona". Necesita ver un caso donde funcionó.

El tercer bloque es el cierre. Diez a trece segundos. Aquí es donde la mayoría de los emprendedores fallan. Terminan con "bueno, eso es todo" o "espero que les haya servido" o simplemente dejan la frase caer. Eso no es un cierre. Es una despedida.

Un cierre tiene dos elementos. El puente: una frase que conecta lo que dijiste con la acción que quieres que tomen. "Si eso que describí te suena familiar, hay algo que puedes hacer ahora mismo." Y la instrucción clara: exactamente qué quieres que hagan. Una acción. Un paso. Sin ambigüedad. "Escríbeme la palabra presencia por mensaje directo y te mando el protocolo." Eso es claro. "Si te interesa búscame" no lo es.

Mira, te voy a advertir del error más frecuente que veo cuando alguien arma su pitch por primera vez. Meten demasiada información en el desarrollo. Quieren explicar todo su método, toda su historia, todos sus resultados. Y el pitch se convierte en un monólogo de tres minutos donde la audiencia se pierde en el segundo cuarenta. Tu pitch no tiene que vender. Tu pitch tiene que abrir una puerta. Tiene que hacer que la persona piense quiero saber más. Si logra eso, cumplió. La venta viene después, en otro formato.

Ahora tu ejercicio. Escribe un pitch de sesenta segundos usando los tres bloques. No para un producto genérico. Para tu oferta actual.

Escribe tu apertura, siete segundos. Elige una herramienta. Léela en voz alta. ¿Tú te detendrías si la vieras en una red social? Si no, reescríbela.

Escribe tu desarrollo, cuarenta segundos. Problema con las palabras de tu cliente, tu enfoque, un resultado concreto. Cronometra. Si pasa de cuarenta y cinco, corta.

Escribe tu cierre, trece segundos. Puente más instrucción clara. ¿Se puede hacer en un solo paso?

Graba el pitch completo usando la técnica de los noventa segundos de la lección tres antes de darle al botón. Escúchalo. Compáralo con el audio de la lección uno. Esa diferencia que escuchas es el sistema funcionando.

En la lección cinco vamos a integrar todo. Las tres capas completas: presencia, control interno y estructura persuasiva, funcionando como un sistema que puedas usar cada día sin pensar. Vamos a asegurarnos de que estas herramientas no se queden en tu cabeza. Que pasen a tu rutina.

Con eso listo, nos vemos en la lección cinco.

Soy Marco Suniaga, y esto es Habla Para Que Actúen."""


def generar_docx(output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    titulo = doc.add_heading("El orden en que dices las cosas lo cambia todo", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Habla Para Que Actúen  |  Lección 4")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    palabras = len(GUION.split())
    duracion = round(palabras / 150)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pairs = [
        ("Duración: ", f"~{duracion} min"),
        ("   Palabras: ", str(palabras)),
        ("   Fecha: ", datetime.now().strftime("%d-%m-%Y")),
        ("   Versión: ", "Humanizada V2"),
    ]
    for label, value in pairs:
        r_label = info.add_run(label)
        r_label.font.size = Pt(10)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(100, 100, 100)
        r_value = info.add_run(value)
        r_value.font.size = Pt(10)
        r_value.font.color.rgb = RGBColor(100, 100, 100)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)

    doc.add_paragraph()

    paragraphs = [p.strip() for p in GUION.split("\n\n") if p.strip()]
    for parrafo in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(40, 40, 40)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep2 = sep2.add_run("─" * 60)
    run_sep2.font.color.rgb = RGBColor(200, 200, 200)
    run_sep2.font.size = Pt(8)
    doc.add_paragraph()

    notas_titulo = doc.add_heading("Notas de grabación", level=2)
    for run in notas_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    notas = [
        "Energía: Pedagógica, clara, con ritmo de enseñanza paso a paso.",
        "Anécdota del colega: Tono de revelación, como si compartiera un secreto aprendido.",
        "Los 3 cerebros: Ritmo constante, explicativo pero no académico.",
        "Ejemplos de apertura: Cambiar ligeramente la voz al citar cada ejemplo, como demostrando.",
        "Framework de 3 bloques: Tono de coach dando instrucciones prácticas.",
        "Ejercicio final: Energía alta, motivacional, empujar a la acción.",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nota)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    anot_titulo = doc.add_heading("Espacio para anotaciones", level=2)
    for run in anot_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    for _ in range(5):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Validado por TrueGuiones  |  Listo para grabar")
    run_footer.font.size = Pt(9)
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return palabras


if __name__ == "__main__":
    out = "/tmp/claude-0/-home-user-Marco/4317bd10-8d2d-5470-84b8-760b686bf996/scratchpad/Leccion-4-El-orden-en-que-dices-las-cosas.docx"
    palabras = generar_docx(out)
    print(f"Generado: {out}")
    print(f"Palabras: {palabras}")
