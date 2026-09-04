#!/usr/bin/env python3
"""Genera el .docx de la Lección 3 — Tu cuerpo habla antes que tu boca"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

GUION = """En las dos lecciones anteriores hicimos un trabajo que la mayoría de los cursos de comunicación se salta. Primero identificaste dónde se quiebra tu mensaje. Después le pusiste nombre a lo que te bloquea. Ahora ya sabes qué disparador se activa en tu cabeza antes de hablar. Hoy vamos a hacer algo diferente. Hoy vas a trabajar con tu cuerpo.

Esto no es una clase teórica sobre lenguaje corporal. Vas a hacer cosas mientras escuchas. Y cuando al final de esta lección te grabes de nuevo, vas a sonar distinto. No porque hayas cambiado de personalidad en veinte minutos. Porque vas a usar herramientas que cambian cómo te perciben.

Te cuento de dónde salen estas herramientas. En televisión, antes de cada escena, yo tenía un ritual. Duraba menos de un minuto. Primero me chequeaba los pies: peso distribuido, nada de balancearme. Después los hombros: atrás y abajo, sin tensión. El mentón paralelo al piso. Y finalmente una respiración lenta antes de que el director dijera acción. Cinco segundos. Ese ritual separaba una escena creíble de una vacía.

Una vez, grabando una escena de confrontación, el director me paró después de dos tomas y me dijo: "Marco, tu texto está bien, pero te estás comiendo el final de cada frase." Yo no me había dado cuenta. Empezaba las frases con fuerza y al llegar a las últimas tres palabras bajaba el volumen, como si mi cerebro ya estuviera en la frase siguiente. Me hizo repetir la escena con una sola indicación: mantén el volumen hasta el punto. Mismas palabras, misma emoción. La toma fue completamente diferente. Porque cuando bajas el volumen al final, comunicas que lo que acabas de decir no importa tanto. Y eso es exactamente lo que hacen la mayoría de los emprendedores cuando hablan.

Mira, tu voz comunica mucho más de lo que crees. No hablo del contenido, hablo del sonido. El tono, el volumen, el ritmo y las pausas. Esas cuatro variables las puedes controlar de forma consciente. Y cuando las controlas, la percepción que generas cambia por completo.

Empecemos por el tono. Cuando estás nervioso, la voz se agudiza porque la tensión aprieta las cuerdas vocales. No necesitas tener voz de locutor. Necesitas que tu tono natural salga, no la versión tensa. Hay un ejercicio que yo hacía antes de cada escena: tararear durante diez segundos y sentir dónde vibra la voz. Si vibra en la garganta o la nariz, el tono sube. Si vibra en el pecho, baja naturalmente. Tu objetivo es llevar la resonancia al pecho. Tararea ahora y lleva la vibración hacia abajo. Nota la diferencia.

El volumen es lo que el director me corrigió en aquella escena. No te estoy pidiendo que grites. Te pido que no te tragues las últimas palabras de cada frase. Empiezas fuerte y vas bajando porque ya estás pensando en lo siguiente. Cuando eso pasa, comunicas inseguridad en lo que acabas de decir. La corrección es simple: mantén el volumen estable hasta el punto final. Ese solo ajuste cambia radicalmente cómo se percibe tu seguridad.

El ritmo. Hablar rápido no es energía. Es nerviosismo disfrazado de dinamismo. Y hablar lento sin variación es monotonía. Lo que necesitas es variación intencional. Aceleras cuando das contexto, datos, antecedentes. Y reduces la velocidad cuando llegas a la idea que quieres que se quede. Si dices "tu producto no es el problema" a la misma velocidad que el contexto previo, pierdes el golpe. La frase de impacto necesita aire. Necesita desaceleración.

Y la pausa. Esta es la herramienta que menos usan los emprendedores y la que más impacto tiene. El silencio intencional. Cuando haces una pausa antes de una frase importante, creas anticipación. El cerebro de quien escucha dice viene algo. Y cuando finalmente hablas, aterriza con más peso. Cuando haces la pausa después, le das tiempo al cerebro para procesar. Sin eso, tu siguiente frase aplasta la anterior y ninguna se queda. Los emprendedores le tienen miedo al silencio. Sienten que si no están hablando, pierden a la audiencia. Es al revés. En televisión, las escenas más poderosas no son las que tienen más diálogo. Son las que tienen más silencio entre diálogos.

Ahora pasemos al cuerpo. Tu cuerpo habla antes que tu boca. Cuando alguien te ve en un video o en una reunión, lo primero que procesa es tu postura, tu expresión, la posición de tus manos. Hay tres ajustes posturales que cambias en dos segundos y que transforman la percepción.

Pies firmes. Peso distribuido en los dos pies. Nada de balancearte ni cambiar de un pie a otro. Si estás sentado, las dos plantas en el piso. Los pies firmes son la base de todo lo demás.

Hombros atrás y abajo. No como un militar. Como alguien que no tiene prisa. Cuando los hombros están arriba y hacia adelante, que es la postura de estrés, comunicas tensión. Cuando están atrás y relajados, comunicas control. Es un ajuste de dos centímetros que cambia la percepción completa.

Mentón paralelo al piso. Ni levantado como si miraras por encima de la gente, ni abajo como si buscaras aprobación. Paralelo. La posición que comunica estoy a tu nivel. En video, esto importa todavía más porque la cámara amplifica cualquier inclinación.

Sobre las manos: visibles y sueltas. No cruzadas, no en los bolsillos. Los gestos naturales aparecen solos cuando dejas de pensar en qué hacer con las manos.

Fíjate que todo lo que te acabo de dar son ajustes mecánicos. Pies, hombros, mentón, resonancia, volumen. No te estoy pidiendo que cambies tu personalidad ni que finjas seguridad. Son herramientas. Las calibras y funcionan independientemente de cómo te sientas por dentro. Eso fue lo que aprendí en un set de televisión. Había días en que no tenía ganas de grabar, días en que estaba preocupado por algo personal. Pero el ritual funcionaba igual. Porque la presencia no depende de tu estado emocional. Depende de tu técnica.

Pero falta un elemento. Puedes tener la voz calibrada y la postura alineada, y aún así tu comunicación puede salir correcta pero plana. Lo que falta es intención. Saber exactamente qué quieres que la persona sienta o piense cuando termines de hablar. No qué quieres decir. Qué quieres lograr.

Te lo digo directo: cuando un actor entra a una escena, no piensa voy a decir estas líneas. Piensa en esta escena mi objetivo es que el otro personaje se sienta atrapado. Esa intención cambia todo. El tono, el ritmo, las pausas, la postura. Todo se alinea hacia el mismo punto.

Cuando un emprendedor graba un video pensando voy a explicar mi producto, la comunicación sale informativa. Correcta, pero plana. Cuando graba pensando quiero que quien vea esto sienta que yo entiendo su problema mejor que él mismo, la comunicación cambia. Se vuelve directa, personal, con peso.

Una emprendedora que trabaja conmigo tenía este problema. Sabía su contenido, hablaba con buena postura, la voz no le temblaba. Pero sus videos no conectaban. Cuando le pregunté qué pensaba antes de grabar, me dijo: "Pienso en no olvidarme de nada." Le cambié la pregunta: ¿qué quieres que la persona sienta cuando termine tu video? Me respondió: "Que por fin alguien la entiende." Grabó el mismo video con esa frase en la cabeza. El contenido era prácticamente igual. La percepción fue otra.

Antes de cualquier comunicación importante, hazte esa pregunta: ¿qué quiero que esta persona sienta o piense cuando yo termine? Si no puedes responderla en una frase, todavía no estás listo para hablar.

Ahora te voy a dar un protocolo que junta todo. Se llama la técnica de los noventa segundos.

Los primeros treinta son el reset físico. Ponte de pie. Pies firmes. Hombros atrás y abajo. Mentón paralelo. Manos sueltas. Respira: cuatro segundos por la nariz, retén cuatro, exhala seis por la boca. Repite tres veces. Esto reduce la respuesta de estrés que vimos en la lección dos.

Los siguientes treinta son el reset vocal. Tararea quince segundos llevando la resonancia al pecho. Después di en voz alta una frase cualquiera, puede ser tu nombre y qué haces, manteniendo el tono bajo, el volumen estable y el ritmo controlado. No estás ensayando. Estás calibrando tu instrumento.

Los últimos treinta son la intención. Responde mentalmente: ¿qué quiero que esta persona sienta o piense cuando yo termine? Una frase. Clara. Concreta. Y con esa frase en la cabeza, empieza.

Noventa segundos. Un protocolo que cualquiera puede hacer antes de darle al botón de grabar o de entrar a una reunión.

Ahora tu ejercicio. Haz la técnica de los noventa segundos completa. Después graba un nuevo audio de sesenta segundos diciendo lo mismo que grabaste en la lección uno. Qué haces y por qué alguien debería trabajar contigo.

Escucha los dos audios seguidos. El de la lección uno y el de ahora. No me tienes que creer. Solo escucha la diferencia. Si no la escuchas, algo del protocolo no lo hiciste completo. Si la escuchas, y la vas a escuchar, tienes la prueba de que esto funciona en tu propia voz.

Ya tienes el diagnóstico, el nombre de tu bloqueo y las herramientas de presencia. Lo que falta es el mensaje. Porque puedes tener toda la presencia del mundo, pero si lo que dices no tiene estructura, no activa decisiones. En la lección cuatro vas a aprender cómo el cerebro de tu audiencia procesa un mensaje persuasivo y vas a construir tu propio pitch con un framework de tres bloques.

Guarda los dos audios. Los vas a necesitar en la lección cinco.

Con eso listo, nos vemos en la lección cuatro.

Soy Marco Suniaga, y esto es Habla Para Que Actúen."""


def generar_docx(output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    titulo = doc.add_heading("Tu cuerpo habla antes que tu boca", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Habla Para Que Actúen  |  Lección 3")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    palabras = len(GUION.split())
    duracion = round(palabras / 150)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pairs = [
        ("Duración: ", f"~{duracion} min"),
        ("   Palabras: ", str(palabras)),
        ("   Fecha: ", datetime.now().strftime("%d-%m-%Y")),
        ("   Versión: ", "Humanizada V2"),
    ]
    for label, value in pairs:
        r_label = info.add_run(label)
        r_label.font.size = Pt(10)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(100, 100, 100)
        r_value = info.add_run(value)
        r_value.font.size = Pt(10)
        r_value.font.color.rgb = RGBColor(100, 100, 100)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)

    doc.add_paragraph()

    paragraphs = [p.strip() for p in GUION.split("\n\n") if p.strip()]
    for parrafo in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(40, 40, 40)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep2 = sep2.add_run("─" * 60)
    run_sep2.font.color.rgb = RGBColor(200, 200, 200)
    run_sep2.font.size = Pt(8)
    doc.add_paragraph()

    notas_titulo = doc.add_heading("Notas de grabación", level=2)
    for run in notas_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    notas = [
        "Energía: Práctica, demostrativa. Más coach que las lecciones anteriores.",
        "Ritmo: Pausar antes y después de cada ejercicio práctico (tarareo, postura).",
        "Anécdota del director: Tono de memoria vívida, como si lo reviviera.",
        "Sección de voz: Ralentizar al dar indicaciones de ejercicio, como guiando.",
        "Sección de cuerpo: Tono firme, instruccional, sin dramatismo.",
        "Emprendedora: Contar con calidez, pausar en la revelación del cambio.",
        "Técnica 90 segundos: Ligeramente más lento, ritmo de instrucción paso a paso.",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nota)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    anot_titulo = doc.add_heading("Espacio para anotaciones", level=2)
    for run in anot_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    for _ in range(5):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Validado por TrueGuiones  |  Listo para grabar")
    run_footer.font.size = Pt(9)
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return palabras


if __name__ == "__main__":
    out = "/tmp/claude-0/-home-user-Marco/4317bd10-8d2d-5470-84b8-760b686bf996/scratchpad/Leccion-3-Tu-cuerpo-habla-antes-que-tu-boca.docx"
    palabras = generar_docx(out)
    print(f"Generado: {out}")
    print(f"Palabras: {palabras}")
