#!/usr/bin/env python3
"""Genera el .docx de la Lección 1 — Dónde se rompe tu mensaje"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import datetime
import os

GUION = """Bienvenido. Me llamo Marco Suniaga y esto es Habla Para Que Actúen.

Antes de meternos en materia te voy a pedir algo. Agarra tu teléfono. Abre la grabadora de voz. Y graba un audio de sesenta segundos explicando qué haces y por qué alguien debería trabajar contigo. Sin preparar nada. Sin pensar demasiado. Sesenta segundos y dale.

¿Lo hiciste? Si lo hiciste, bien. Si no, pausa esto y hazlo. Todo lo que viene después se construye sobre ese audio. Y si te saltaste el ejercicio pensando que después lo haces, te lo digo directo: no funciona así. Necesito que lo grabes ahora porque vamos a trabajar con ese material en los próximos minutos.

Ahora escúchalo. Y mientras lo escuchas quiero que te fijes en tres cosas.

Primera: ¿dónde dudaste? Hay un momento en tu audio donde tu voz tembló, hiciste una pausa rara, dijiste "eh" o "bueno" para ganar tiempo. Ese momento te está diciendo exactamente dónde tu mensaje se quiebra. Qué parte de tu oferta todavía no la tienes resuelta en tu cabeza. Porque cuando algo lo tienes claro, las palabras salen. Cuando algo todavía está a medio cocinar, tu voz lo delata. Siempre.

Segunda: ¿dónde aceleraste? Cuando alguien acelera es porque quiere pasar rápido por algo que le incomoda. A lo mejor es cuando mencionas el precio. O cuando llegas a la promesa de resultado y te da cosa sonar exagerado. O esa parte donde tienes que decir algo que sientes presuntuoso. Fíjate qué estabas diciendo en ese punto. Probablemente es donde menos seguro te sientes de lo que ofreces.

Tercera: ¿dónde sentiste que sonabas falso? Esta es la más difícil porque es subjetiva. Pero tú lo sabes. Hay una parte donde tú mismo escuchas y piensas "eso no suena como yo quiero sonar". Tal vez usaste un tono que no es tuyo. Tal vez repetiste una frase que escuchaste en algún video de marketing. Tal vez intentaste sonar más seguro de lo que te sentías. Anota qué estabas diciendo ahí.

Esos tres puntos, la duda, la aceleración y la falsedad, son tu diagnóstico. Y te digo algo: la mayoría de los emprendedores que llegan a mí nunca se han escuchado hablar. Nunca. Graban videos, hacen llamadas de venta, presentaciones, webinars. Pero nunca se sientan a escucharse con atención. Entonces repiten los mismos errores semana tras semana y se preguntan por qué la gente escucha pero no compra.

Mira, yo vengo del mundo del entretenimiento. Trabajé en televisión, en telenovelas, en cine. Y hay algo que aprendí en esos años que aplica directamente a lo que tú haces como emprendedor. En un set de televisión tienes muy poco tiempo para transmitir algo real. A veces una escena dura cuarenta segundos. Cuarenta segundos para que el espectador crea que eres ese personaje, sienta lo que el personaje siente y no cambie de canal.

Recuerdo una escena que grabé donde mi personaje tenía que convencer a otro de tomar una decisión arriesgada. El director me paró y me dijo: "Marco, las palabras están bien. El texto está perfecto. Pero estás desconectado de lo que dices. Estás recitando." Y me hizo repetir la escena cambiando solamente una cosa: bajar la velocidad en una frase y hacer contacto visual antes de la palabra más importante. Mismas palabras. Resultado completamente diferente.

Eso es técnica. Aprendes a controlar el tono de tu voz para que una frase suene a autoridad o a cercanía según lo que necesites. Aprendes a usar el cuerpo para que tu postura comunique algo antes de que abras la boca. Y aprendes a manejar los silencios. Un silencio de dos segundos bien puesto hace que lo siguiente que digas tenga tres veces más peso.

¿Y sabes dónde funciona esto exactamente igual? En una llamada de Zoom. En un video para redes. En un webinar de ventas. En una conversación de cierre con un cliente. Tu cerebro no distingue entre un actor profesional y un emprendedor que domina su comunicación. Lo que distingue es si la persona que habla transmite convicción o transmite duda.

Te pongo un ejemplo. He trabajado con emprendedores que tenían cursos y servicios excelentes. Contenido que de verdad transformaba a sus clientes, gente que daba resultados reales. Pero cuando les tocaba grabar un video de ventas o explicar qué hacían en una llamada, sonaban como cualquier otro. Las mismas frases que repiten todos. El mismo tono neutro. El cliente potencial escuchaba treinta segundos, pensaba "esto ya lo escuché" y cerraba la pestaña.

Una emprendedora con la que trabajé el año pasado tenía un programa de nutrición para madres emprendedoras. Programa bueno, con testimonios reales. Pero cada vez que grababa un video para promocionarlo, hablaba como si estuviera leyendo un guión que alguien le escribió. Le pregunté "¿tú hablas así cuando le explicas a una amiga lo que haces?" Me dijo que no. "Entonces ¿por qué hablas así en cámara?" Y ahí cayó la ficha. Tenía un modo de comunicación para su vida real y otro modo para "vender". Y el modo de vender era genérico, prestado, sin personalidad.

Eso pasa más de lo que crees. Y la consecuencia es seria: puedes tener el mejor producto del mercado, pero si cuando hablas suenas igual que los otros diez que ofrecen algo parecido, el cliente potencial no tiene forma de diferenciarte. Compara precio y se va con el más barato.

Ahora, yo podría decirte "practica más" y dejarte ahí. Pero hay un problema real con practicar sin dirección. La práctica refuerza lo que ya haces. Si lo que haces está desalineado, practicar más te hace más eficiente en algo desalineado. Es como ir al gimnasio un año con mala forma en los ejercicios. Vas a ver algún resultado, claro. Pero estás lejos de lo que lograrías con la técnica correcta. Y con el tiempo puedes lesionarte.

En comunicación pasa lo mismo. Puedes grabar cien videos y seguir sonando genérico si nadie te señala dónde se quiebra tu mensaje y por qué.

Por eso este mini curso tiene un sistema. Tres capas que yo uso desde hace años y que combinan lo que aprendí en escenario con dos disciplinas que le dan estructura a todo.

La primera capa es Presencia. Tu voz, tu cuerpo, tu energía. Lo que la gente percibe antes de procesar tus palabras. Viene directo de la técnica de actor. Son habilidades que se entrenan con ejercicios concretos. Y las vamos a entrenar.

La segunda es Control interno. Lo que pasa dentro de tu cabeza antes y durante una comunicación que te importa. Los bloqueos, esas voces internas que te dicen "vas a hacer el ridículo" o "¿quién eres tú para enseñar esto?" Viene de la programación neurolingüística. Son protocolos específicos para identificar qué te frena y desactivarlo. Porque de nada sirve que yo te enseñe herramientas de comunicación si cada vez que te toca usarlas tu cabeza te sabotea y terminas cancelando el video o improvisando mal la llamada.

La tercera es Estructura persuasiva. El orden en que presentas tu mensaje para que active decisiones en el cerebro de quien te escucha. Viene del neuromarketing. Tu cerebro procesa información en un orden específico. Si respetas ese orden, tu mensaje llega con más fuerza. Si lo violas, tu mensaje se pierde aunque el contenido sea brillante. Y la mayoría de los emprendedores violan ese orden sin saberlo, porque nadie les enseñó cómo funciona la toma de decisiones desde la perspectiva del que escucha.

Esas tres capas forman lo que vas a aprender en estas cinco lecciones. Esta primera fue diagnóstico, que sepas dónde estás parado. La segunda va sobre tu cabeza, sobre qué hace tu cerebro cuando te bloqueas y cómo identificar tu disparador principal. Tu disparador es la situación específica que activa tu bloqueo, algo mucho más preciso que "los nervios", y en la lección dos le vamos a poner nombre. La tercera es presencia física, cómo cambiar la percepción que generas cuando hablas. La cuarta es estructura, cómo ordenar tu mensaje para que active decisiones. Y la quinta integra todo en un sistema que puedas usar cada día sin pensar.

Y te lo digo directo: mi estándar para este curso es que si al final de las cinco lecciones sabes más de oratoria pero sigues sin poder cerrar una venta con tu voz, fallé yo. Esto va de que la persona que te escucha actúe. Compre, se inscriba, diga que sí, cambie de opinión. Eso.

Antes de pasar a la lección dos necesito que hagas una cosa más. Escribe una frase. Una sola. Completa esto: "Mi mayor problema al comunicar es..." y la terminas con tus palabras. Escríbela y guárdala. Esa frase la vas a necesitar en la lección dos.

Asegúrate de tener tu audio de sesenta segundos grabado, las tres anotaciones hechas y tu frase escrita. Con eso nos vemos en la lección dos.

Soy Marco Suniaga, y esto es Habla Para Que Actúen."""


def generar_docx(output_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # === TITLE ===
    titulo = doc.add_heading("Dónde se rompe tu mensaje", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(26)

    # === SUBTITLE ===
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Habla Para Que Actúen  |  Lección 1")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    # === METADATA BAR ===
    palabras = len(GUION.split())
    duracion = round(palabras / 150)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pairs = [
        ("Duración: ", f"~{duracion} min"),
        ("   Palabras: ", str(palabras)),
        ("   Fecha: ", datetime.now().strftime("%d-%m-%Y")),
        ("   Versión: ", "Humanizada V2"),
    ]
    for label, value in pairs:
        r_label = info.add_run(label)
        r_label.font.size = Pt(10)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(100, 100, 100)
        r_value = info.add_run(value)
        r_value.font.size = Pt(10)
        r_value.font.color.rgb = RGBColor(100, 100, 100)

    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)

    doc.add_paragraph()

    # === SCRIPT BODY ===
    paragraphs = [p.strip() for p in GUION.split("\n\n") if p.strip()]

    for parrafo in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(40, 40, 40)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # === SEPARATOR ===
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep2 = sep2.add_run("─" * 60)
    run_sep2.font.color.rgb = RGBColor(200, 200, 200)
    run_sep2.font.size = Pt(8)

    doc.add_paragraph()

    # === RECORDING NOTES ===
    notas_titulo = doc.add_heading("Notas de grabación", level=2)
    for run in notas_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    notas = [
        "Energía: Alta, directa, con autoridad natural. Sin dramatismo.",
        "Ritmo: Mezcla agresiva de frases muy cortas con desarrollos largos.",
        "Pausas: Después de preguntas directas y antes de revelaciones.",
        "Marcadores: \"Mira\", \"Te pongo un ejemplo\", \"Fíjate\", \"Te lo digo directo\".",
        "Oralidad: Suena a conversación uno a uno, no a conferencia.",
        "Apertura: Arranca con acción (ejercicio del audio), no con bienvenida genérica.",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nota)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.space_after = Pt(4)

    # === ANNOTATIONS SPACE ===
    doc.add_paragraph()
    doc.add_paragraph()

    anot_titulo = doc.add_heading("Espacio para anotaciones", level=2)
    for run in anot_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    for _ in range(5):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    # === FOOTER ===
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Validado por TrueGuiones  |  Listo para grabar")
    run_footer.font.size = Pt(9)
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return palabras


if __name__ == "__main__":
    out = "/tmp/claude-0/-home-user-Marco/4317bd10-8d2d-5470-84b8-760b686bf996/scratchpad/Leccion-1-Donde-se-rompe-tu-mensaje.docx"
    palabras = generar_docx(out)
    print(f"Generado: {out}")
    print(f"Palabras: {palabras}")
