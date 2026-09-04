#!/usr/bin/env python3
"""Genera el .docx de la Lección 5 — Tu sistema completo, de saber a ejecutar"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

GUION = """Llegaste a la lección cinco. La última. Y te voy a decir algo que puede sonar contradictorio: esta lección no tiene contenido nuevo. Todo lo que necesitas ya lo tienes. Lo que vamos a hacer hoy es algo que la mayoría de los cursos no hace. Asegurarnos de que lo que aprendiste no se quede en tu cabeza como información. Que se convierta en algo que haces.

Porque hay una diferencia enorme entre saber y hacer. Saber es que la postura afecta cómo te perciben. Hacer es pararte frente a una cámara con los pies firmes, los hombros atrás, el mentón paralelo, y que eso salga automático. Saber es que tu pitch necesita tres bloques. Hacer es abrir la boca en una reunión y que esos tres bloques salgan en orden, naturales, sin que nadie note que hay una estructura detrás.

Te cuento algo que me pasó con un cliente hace unos meses. Tipo inteligente, había tomado tres cursos de comunicación antes del mío. Sabía de lenguaje corporal, sabía de storytelling, sabía de neuroventas. Me recitaba conceptos como un profesor. Pero cuando le pedí que grabara un video de sesenta segundos presentando su servicio, se trabó en los primeros diez. Le pregunté qué había pasado. Me dijo: "Es que sé demasiado y no sé por dónde empezar." Tenía conocimiento sin sistema. Piezas sueltas sin un orden de ejecución. Lo pusimos a usar el protocolo de tres minutos, le di la estructura de tres bloques, y en la segunda semana había grabado cuatro videos y cerrado una venta por llamada. Organizó lo que ya sabía y eso fue suficiente.

Vamos a repasar lo que tienes. A lo largo de cuatro lecciones construiste tres capas. La primera es presencia. La trabajaste en la lección tres. Cuatro variables de voz: tono, volumen, ritmo, pausa. Tres ajustes posturales: pies, hombros, mentón. Herramientas mecánicas que cambian cómo te perciben antes de que procesen tus palabras. No dependen de cómo te sientas. Las aplicas y funcionan.

La segunda es control interno. Lección dos. Identificaste tu disparador principal. Entendiste el patrón: detonante, pensamiento automático, respuesta. Y aprendiste que la intervención real no está en respirar hondo o calmarte. Está en observar el pensamiento automático sin obedecerlo.

La tercera es estructura persuasiva. Lección cuatro. Tu pitch de tres bloques diseñado para activar los tres cerebros en el orden correcto. Reptiliano para la atención, límbico para la conexión, neocórtex para la justificación.

Cada capa resuelve un problema diferente. Si la gente no te percibe con autoridad, trabajas presencia. Si te bloqueas antes de hablar, trabajas control interno. Si hablas pero no genera acción, trabajas estructura. La mayoría de los emprendedores tiene uno de estos problemas resuelto a medias. Tal vez se ven bien en cámara pero su mensaje no genera acción. O su mensaje es sólido pero se bloquean antes de entregarlo. El sistema funciona cuando las tres capas están activas al mismo tiempo.

Ahora te voy a dar el protocolo completo. El que junta todo. Menos de tres minutos antes de cualquier comunicación que importe. Una presentación, un video, una llamada de ventas, un audio para un cliente. Esos tres minutos hacen la diferencia entre comunicar al sesenta por ciento de tu capacidad y comunicar al noventa.

El primer minuto es el reset físico y vocal. Ponte de pie. Pies firmes. Hombros atrás y abajo. Mentón paralelo. Manos sueltas. Respiración: cuatro segundos por la nariz, retén cuatro, exhala seis por la boca. Tres ciclos. Tararea quince segundos llevando la vibración al pecho. Di una frase corta en voz alta manteniendo el tono estable y el volumen constante hasta el punto final. Tu cuerpo y tu voz están calibrados.

El segundo minuto es el chequeo de disparador. Pregúntate qué pensamiento automático está apareciendo ahora mismo. Si no hay ninguno, avanza. Si aparece uno, van a pensar que no sé, esto no va a salir bien, de qué sirve si nunca funciona, haz esto: reconócelo. Ese pensamiento apareció. Lo veo. No necesito obedecerlo para hablar. No lo pelees. No lo contradigas. Solo obsérvalo y decide no actuar según él.

El tercer minuto es intención y estructura. Responde mentalmente: ¿qué quiero que esta persona sienta o piense cuando yo termine? Repasa tu estructura de tres bloques. ¿Cuál es tu apertura? ¿Cuál es el problema que vas a nombrar? ¿Cuál es tu cierre y qué acción pides? No necesitas repasarlo palabra por palabra. Solo el esqueleto. Los detalles salen solos cuando la estructura está firme.

Tres minutos. Tres capas. Y entras a la comunicación desde un lugar completamente diferente al que entrarías si simplemente le dieras al botón de grabar.

Mira, ahora tengo que ser directo contigo sobre algo. Porque si no te lo digo, este curso habría fallado en lo más importante.

Todo lo que te enseñé funciona. Las herramientas están probadas. Yo las uso. Mis clientes las usan. Pero nada de esto funciona si no lo aplicas de forma consistente. Y la consistencia es exactamente donde la mayoría se cae.

¿Sabes qué pasa normalmente? La persona termina el curso motivada. Aplica las herramientas tres o cuatro días. Siente la diferencia. Le gusta. Y después la vida pasa. Tiene un día complicado, no hace el protocolo, graba un video sin prepararse, el video no le gusta, y poco a poco vuelve a sus patrones anteriores. No porque las herramientas dejaron de funcionar. Porque dejó de usarlas.

Eso no es un defecto tuyo. El cerebro prefiere lo conocido, aunque lo conocido sea menos efectivo. Cambiar un patrón requiere repetición sostenida. No motivación. No inspiración. Repetición.

Te lo digo directo: en cinco lecciones puedo darte las herramientas y demostrarte que funcionan. Pero no puedo acompañarte durante las semanas que toma convertir esto en un hábito. Eso requiere otro formato, otra estructura, otro nivel de acompañamiento.

Dicho eso, vamos a hacer lo máximo posible con lo que tienes. Te doy un plan de acción para los próximos siete días. Concreto y realista.

Día uno: elige una situación real de esta semana donde necesites comunicar algo. Una reunión, un video, una llamada. Algo que ya tengas agendado o que puedas agendar.

Día dos: prepara tu intervención con el protocolo completo. Los tres minutos. Después escribe tu pitch de tres bloques para esa situación específica. No uses el pitch genérico de la lección cuatro. Adapta la estructura al contexto real.

Día tres: ejecuta. Haz la comunicación. Usa el protocolo antes, la estructura durante. Después anota qué funcionó, qué no, en qué momento perdiste la estructura, en qué momento volviste a tu patrón anterior.

Día cuatro: ajusta. Revisa tus notas. ¿La apertura funcionó? ¿El cierre fue claro? ¿Hiciste el reset antes o se te olvidó? Corrige lo que falló.

Día cinco: repite con una segunda situación. Diferente contexto, mismo protocolo.

Día seis: repite con una tercera. Tres ejecuciones reales en una semana es lo mínimo para empezar a crear el patrón nuevo.

Día siete: evalúa. Escucha tus audios o recuerda tus intervenciones. Ahora escucha el audio de la lección uno, el primer diagnóstico. Compara. No lo que sabes. Lo que haces. ¿Se nota la diferencia? ¿Dónde mejoró? ¿Dónde falta?

Escribe dos frases: antes de este curso, mi mayor problema al comunicar era... Después de este curso, lo que cambió fue... Esas dos frases son tu medida real de progreso. No las mías. Las tuyas.

Fíjate en algo. Si hiciste los ejercicios y seguiste el plan, probablemente también notaste que hay un trecho entre donde estás ahora y donde podrías estar con más práctica guiada, más retroalimentación y más estructura. Esa es la diferencia entre un mini curso y un programa completo. Un mini curso te da el sistema. Un programa te da el acompañamiento para dominarlo.

Yo tengo una masterclass gratuita donde profundizo en exactamente eso. No es una repetición de lo que viste aquí. Es el siguiente paso. Cómo se aplica el sistema de tres capas en escenarios reales de ventas, presentaciones y contenido, con ejemplos grabados. El error más caro que cometen los emprendedores que ya aprendieron a comunicar bien, y que probablemente tú vas a cometer si no lo ves a tiempo. Y el modelo completo de comunicación persuasiva que va más allá de un pitch de sesenta segundos.

Si completaste las cinco lecciones, estás en el punto exacto donde la masterclass tiene más sentido. No antes, porque necesitabas este contexto. No después, porque si esperas mucho, lo que aprendiste se enfría. El enlace está debajo de este video. Es gratuita y dura aproximadamente una hora.

Te voy a cerrar con algo que creo de verdad. Tú ya tienes algo que vale la pena comunicar. Un producto, un servicio, un conocimiento que puede cambiar la situación de alguien. Lo que te faltaba no era más contenido ni más confianza. Lo que te faltaba eran herramientas.

Ahora las tienes. Presencia. Control interno. Estructura persuasiva. Un protocolo de tres minutos. Un framework de tres bloques. Un plan de siete días.

La próxima vez que tengas que comunicar algo importante, en lugar de improvisar, saca el protocolo. Haz el reset. Define tu intención. Y habla con la certeza de que lo que dices va a llegar de la forma en que quieres que llegue. No porque seas un comunicador nato. Porque tienes un sistema.

Soy Marco Suniaga, y esto fue Habla Para Que Actúen. Nos vemos en la masterclass."""


def generar_docx(output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    titulo = doc.add_heading("Tu sistema completo, de saber a ejecutar", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Habla Para Que Actúen  |  Lección 5")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()

    palabras = len(GUION.split())
    duracion = round(palabras / 150)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pairs = [
        ("Duración: ", f"~{duracion} min"),
        ("   Palabras: ", str(palabras)),
        ("   Fecha: ", datetime.now().strftime("%d-%m-%Y")),
        ("   Versión: ", "Humanizada V2"),
    ]
    for label, value in pairs:
        r_label = info.add_run(label)
        r_label.font.size = Pt(10)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(100, 100, 100)
        r_value = info.add_run(value)
        r_value.font.size = Pt(10)
        r_value.font.color.rgb = RGBColor(100, 100, 100)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = sep.add_run("─" * 60)
    run_sep.font.color.rgb = RGBColor(200, 200, 200)
    run_sep.font.size = Pt(8)

    doc.add_paragraph()

    paragraphs = [p.strip() for p in GUION.split("\n\n") if p.strip()]
    for parrafo in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(parrafo)
        run.font.size = Pt(12)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(40, 40, 40)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep2 = sep2.add_run("─" * 60)
    run_sep2.font.color.rgb = RGBColor(200, 200, 200)
    run_sep2.font.size = Pt(8)
    doc.add_paragraph()

    notas_titulo = doc.add_heading("Notas de grabación", level=2)
    for run in notas_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    notas = [
        "Energía: Integradora, de cierre. Empezar calmado, subir en el protocolo, bajar en la honestidad.",
        "Anécdota del cliente: Tono de caso real, ritmo narrativo rápido.",
        "Recapitulación 3 capas: Ritmo constante, como checklist, sin dramatismo.",
        "Protocolo completo: Tono de coach guiando paso a paso, pausas entre cada minuto.",
        "Sección de honestidad: Bajar energía, tono íntimo, como confesión profesional.",
        "Plan de 7 días: Energía alta, instruccional, empujar a la acción concreta.",
        "Puente a masterclass: Natural, sin tono de venta, como recomendación de colega.",
        "Cierre final: Energía emotiva contenida, pausar antes de la última frase.",
    ]
    for nota in notas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(nota)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    anot_titulo = doc.add_heading("Espacio para anotaciones", level=2)
    for run in anot_titulo.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
        run.font.size = Pt(14)

    for _ in range(5):
        p = doc.add_paragraph()
        run = p.add_run("_" * 85)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("Validado por TrueGuiones  |  Listo para grabar")
    run_footer.font.size = Pt(9)
    run_footer.font.italic = True
    run_footer.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(output_path)
    return palabras


if __name__ == "__main__":
    out = "/tmp/claude-0/-home-user-Marco/4317bd10-8d2d-5470-84b8-760b686bf996/scratchpad/Leccion-5-Tu-sistema-completo.docx"
    palabras = generar_docx(out)
    print(f"Generado: {out}")
    print(f"Palabras: {palabras}")
