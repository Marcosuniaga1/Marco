#!/usr/bin/env python3
"""Genera propuesta ROI para cliente inmobiliaria — MAS Marketing Agency"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

# ── Colores de marca ──────────────────────────────────────────────────────────
GOLD_HEX   = "C9A42C"
BLACK_HEX  = "1A1A1A"
WHITE_HEX  = "FFFFFF"
GRAY_HEX   = "F4F4F4"
LGRAY_HEX  = "E8E8E8"
DARKGRAY   = "555555"

GOLD  = RGBColor(0xC9, 0xA4, 0x2C)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0xF4, 0xF4, 0xF4)
DGRAY = RGBColor(0x55, 0x55, 0x55)

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_borders(cell, color_hex="C9A42C", size="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font_name="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def section_title(doc, text):
    """Línea dorada con título de sección"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0)
    # Barra lateral dorada via shading de párrafo — usamos run con borde
    add_run(p, "  ", size=12)
    run = add_run(p, f"  {text}", bold=True, size=13, color=BLACK, font_name="Calibri")
    # Subrayado con color dorado
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), GOLD_HEX)
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def build_table(doc, headers, rows, col_widths=None,
                header_bg=BLACK_HEX, header_fg=WHITE,
                stripe_bg=GRAY_HEX, highlight_last=False):
    """Crea tabla con estilo premium"""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'

    # Encabezado
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell_bg(cell, header_bg)
        cell_borders(cell, GOLD_HEX, "6")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        add_run(p, h, bold=True, size=10, color=header_fg, font_name="Calibri")
        if col_widths:
            cell.width = Inches(col_widths[i])

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        is_last = (r_idx == len(rows) - 1)
        bg = (GOLD_HEX if (highlight_last and is_last) else
              (GRAY_HEX if r_idx % 2 == 0 else WHITE_HEX))
        fg = (WHITE if (highlight_last and is_last) else BLACK)

        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell_bg(cell, bg)
            cell_borders(cell, GOLD_HEX if (highlight_last and is_last) else LGRAY_HEX, "4")
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            bold = (is_last and highlight_last) or (c_idx == 0)
            add_run(p, str(val), bold=bold, size=10,
                    color=WHITE if (highlight_last and is_last) else BLACK,
                    font_name="Calibri")
            if col_widths:
                cell.width = Inches(col_widths[c_idx])

    doc.add_paragraph()
    return table

# ── Documento ─────────────────────────────────────────────────────────────────
def generar(output_path):
    doc = Document()

    # Márgenes
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.left_margin  = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.top_margin   = Inches(0.6)
    sec.bottom_margin = Inches(0.75)

    # ── ENCABEZADO NEGRO CON LOGO / MARCA ────────────────────────────────────
    logo_path = os.path.join(os.path.dirname(__file__), "mas_logo.png")

    # Header table: logo izquierda | texto derecha
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = 'Table Grid'
    hdr_row = hdr_table.rows[0]

    # Celda izquierda — logo
    left_cell = hdr_row.cells[0]
    left_cell.width = Inches(1.3)
    cell_bg(left_cell, BLACK_HEX)
    no_border(left_cell)
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(8)
    lp.paragraph_format.space_after  = Pt(8)
    if os.path.exists(logo_path):
        run = lp.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
    else:
        add_run(lp, "MAS", bold=True, size=24, color=GOLD, font_name="Calibri")

    # Celda derecha — nombre y slogan
    right_cell = hdr_row.cells[1]
    right_cell.width = Inches(5.8)
    cell_bg(right_cell, BLACK_HEX)
    no_border(right_cell)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp.paragraph_format.space_before = Pt(14)
    rp.paragraph_format.left_indent  = Inches(0.2)
    add_run(rp, "MAS MARKETING AGENCY", bold=True, size=18, color=GOLD, font_name="Calibri")
    rp2 = right_cell.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp2.paragraph_format.left_indent = Inches(0.2)
    rp2.paragraph_format.space_after = Pt(12)
    add_run(rp2, "Soluciones de Inteligencia Artificial para Negocios", italic=True,
            size=10, color=RGBColor(0xAA, 0xAA, 0xAA), font_name="Calibri")

    doc.add_paragraph()

    # ── TÍTULO PRINCIPAL ─────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(4)
    add_run(title_p, "PROYECCIÓN DE RETORNO DE INVERSIÓN",
            bold=True, size=18, color=BLACK, font_name="Calibri")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(2)
    add_run(sub_p, "Agente de Inteligencia Artificial en WhatsApp para Inmobiliarias",
            size=12, color=DGRAY, font_name="Calibri")

    # Línea separadora dorada
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(14)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), GOLD_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── INTRO ────────────────────────────────────────────────────────────────
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    add_run(intro, "Estimado(a) [Nombre del Cliente],\n", bold=True, size=11, color=BLACK)
    add_run(intro,
        "Antes de tomar una decisión, queremos que evalúes esta propuesta "
        "con números concretos. No se trata de un gasto — se trata de una "
        "inversión con retorno medible desde el primer mes.",
        size=11, color=RGBColor(0x33,0x33,0x33))

    # ══ SECCIÓN 1 ════════════════════════════════════════════════════════════
    section_title(doc, "La realidad de tu inmobiliaria HOY (sin el agente)")

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    add_run(p1,
        "Cada mes recibes decenas de mensajes de personas interesadas en tus propiedades. "
        "Pero la atención manual tiene un límite — y ese límite te está costando dinero.",
        size=10.5, color=RGBColor(0x44,0x44,0x44))

    build_table(doc,
        headers=["INDICADOR", "SITUACIÓN ACTUAL"],
        rows=[
            ["Mensajes de WhatsApp recibidos al mes",     "~50 mensajes"],
            ["Mensajes atendidos a tiempo",               "~28  (56%)"],
            ["Mensajes perdidos (noche, fin de semana, ocupado)", "~22  (44%)"],
            ["Leads que se convierten en visita",         "~7 visitas/mes"],
            ["Visitas que cierran operación",             "1 – 2 cierres/mes"],
            ["Comisión promedio por operación",           "$2.500 – $4.000 USD"],
            ["Ingreso mensual estimado",                  "$2.500 – $8.000 USD"],
        ],
        col_widths=[3.8, 2.8],
        header_bg=BLACK_HEX,
    )

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after  = Pt(12)
    callout.paragraph_format.left_indent  = Inches(0.3)
    add_run(callout,
        "⚠  El 44% de las personas que te contactan nunca reciben respuesta a tiempo. "
        "Cuando alguien tarda en responder en WhatsApp, el cliente simplemente llama a otra inmobiliaria.",
        italic=True, size=10, color=RGBColor(0x99,0x60,0x00))

    # ══ SECCIÓN 2 ════════════════════════════════════════════════════════════
    section_title(doc, "Lo que cambia con el Agente IA")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    add_run(p2,
        "El agente responde en segundos, a las 11pm, un domingo, en feriado. "
        "Califica al comprador, presenta opciones según su presupuesto y agenda la visita — "
        "sin que tú hagas nada.",
        size=10.5, color=RGBColor(0x44,0x44,0x44))

    build_table(doc,
        headers=["INDICADOR", "CON AGENTE IA"],
        rows=[
            ["Mensajes atendidos al mes",                 "~50  (100%)"],
            ["Disponibilidad",                            "24 / 7 / 365"],
            ["Leads calificados que se convierten en visita", "~14 – 18 visitas/mes"],
            ["Visitas que cierran operación",             "2 – 4 cierres/mes"],
            ["Comisión promedio por operación",           "$2.500 – $4.000 USD"],
            ["Ingreso mensual estimado",                  "$5.000 – $16.000 USD"],
        ],
        col_widths=[3.8, 2.8],
        header_bg=BLACK_HEX,
    )

    # ══ SECCIÓN 3 ════════════════════════════════════════════════════════════
    section_title(doc, "La comparativa que importa")

    build_table(doc,
        headers=["INDICADOR", "SIN AGENTE", "CON AGENTE IA", "DIFERENCIA"],
        rows=[
            ["Cierres al mes",          "1 – 2",          "2 – 4",          "+1 a +2 operaciones"],
            ["Ingreso mensual",         "$5.000 USD",      "$10.000 USD",     "+$5.000 USD"],
            ["Costo del agente",        "—",              "$180 / mes",      "—"],
            ["Ganancia adicional neta", "—",              "—",              "+ $4.820 USD / mes"],
        ],
        col_widths=[2.5, 1.5, 1.8, 2.1],
        header_bg=BLACK_HEX,
        highlight_last=True,
    )

    # ══ SECCIÓN 4 ════════════════════════════════════════════════════════════
    section_title(doc, "Tu inversión — lo que realmente pagas")

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    add_run(p3,
        "La implementación se realiza una sola vez. A partir del mes 1 "
        "el agente trabaja de forma autónoma por un costo fijo mensual.",
        size=10.5, color=RGBColor(0x44,0x44,0x44))

    build_table(doc,
        headers=["CONCEPTO", "PRECIO"],
        rows=[
            ["Implementación y personalización completa  (pago único)", "$350 USD"],
            ["Mantenimiento mensual — API + soporte + actualizaciones",  "$180 USD / mes"],
            ["Integración con Google Calendar y Google Sheets",          "Incluido"],
            ["Panel de leads y visitas en tiempo real",                  "Incluido"],
            ["Soporte técnico mensual",                                  "Incluido"],
        ],
        col_widths=[4.2, 2.4],
        header_bg=BLACK_HEX,
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after  = Pt(6)
    add_run(note, "Precio de lanzamiento — cupos limitados a 5 clientes este mes.",
            bold=True, italic=True, size=10, color=GOLD)

    # ══ SECCIÓN 5 — ROI ══════════════════════════════════════════════════════
    section_title(doc, "Tu retorno de inversión mes a mes")

    build_table(doc,
        headers=["PERÍODO", "INVERSIÓN ACUMULADA", "GANANCIA ADICIONAL ESTIMADA", "RETORNO NETO"],
        rows=[
            ["Mes 1  (setup + 1 mensualidad)", "$530 USD",  "+$5.000 USD",  "+$4.470 USD"],
            ["Mes 3",                           "$710 USD",  "+$15.000 USD", "+$14.290 USD"],
            ["Mes 6",                           "$1.250 USD","+$30.000 USD", "+$28.750 USD"],
            ["Mes 12",                          "$2.510 USD","+$60.000 USD", "+$57.490 USD"],
        ],
        col_widths=[2.0, 2.1, 2.3, 1.9],
        header_bg=BLACK_HEX,
        highlight_last=True,
    )

    # ══ CIERRE ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    cierre_p = doc.add_paragraph()
    cierre_p.paragraph_format.space_before = Pt(6)
    cierre_p.paragraph_format.space_after  = Pt(6)
    add_run(cierre_p,
        "La pregunta no es si puedes permitirte el agente.\n",
        bold=True, size=13, color=BLACK)
    add_run(cierre_p,
        "La pregunta es cuánto te está costando ",
        size=13, color=BLACK)
    add_run(cierre_p,
        "no tenerlo.",
        bold=True, size=13, color=GOLD)

    cta = doc.add_paragraph()
    cta.paragraph_format.space_before = Pt(12)
    cta.paragraph_format.space_after  = Pt(4)
    add_run(cta,
        "¿Arrancamos esta semana? Contáctame y en 48 horas tu agente está operativo.",
        bold=True, size=11, color=BLACK)

    # ── FOOTER ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2.paragraph_format.space_before = Pt(10)
    pPr2 = sep2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top2 = OxmlElement('w:top')
    top2.set(qn('w:val'), 'single')
    top2.set(qn('w:sz'), '8')
    top2.set(qn('w:space'), '1')
    top2.set(qn('w:color'), GOLD_HEX)
    pBdr2.append(top2)
    pPr2.append(pBdr2)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(4)
    add_run(footer_p, "MAS MARKETING AGENCY  |  marcosuniagaoficial@gmail.com",
            size=9, color=DGRAY, font_name="Calibri")

    doc.save(output_path)
    print(f"✅  Documento guardado: {output_path}")

# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out = "/home/user/Marco/Propuesta_ROI_MasMarketing.docx"
    generar(out)
